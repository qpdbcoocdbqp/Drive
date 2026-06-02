import os
import pymupdf
import sys
import argparse
from docx import Document

def extract_text(file_path: str) -> str:
    """Extract text from PDF or DOCX file, including tables."""
    if not os.path.exists(file_path):
        return f"Error: File not found - {file_path}"
    
    ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if ext == ".pdf":
            doc = pymupdf.open(file_path)
            full_text = ""
            for page in doc:
                full_text += page.get_text("text") + "\n"
            doc.close()
            return full_text.strip()
            
        elif ext in [".docx", ".doc"]:
            doc = Document(file_path)
            full_text_list = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                full_text_list.append(paragraph.text)
                
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = []
                        for paragraph in cell.paragraphs:
                            cell_text.append(paragraph.text)
                        row_text.append(" ".join(cell_text))
                    full_text_list.append("\n".join(row_text))
            
            return "\n".join(full_text_list).strip()
            
        else:
            return f"Error: Unsupported file format - {ext}"
            
    except Exception as e:
        return f"Error processing document: {str(e)}"

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Extract raw text content from PDF or DOCX documents.")
    parser.add_argument("--file", type=str, required=True, help="Path to the document file (.pdf or .docx).")
    args = parser.parse_args()
    
    file_path = args.file
    result = extract_text(file_path)
    print(result)